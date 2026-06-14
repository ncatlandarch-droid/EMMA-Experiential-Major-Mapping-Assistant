"""Generate FIPSE PSSG Grant Word Documents for Google Drive upload."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.expanduser("~/Desktop")

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '003366')
    
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')
    
    return table

def create_strategy_brief():
    """Create the Strategy Brief Word document."""
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    
    # Title
    title = doc.add_heading('FIPSE PSSG Strategy Brief', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('EMMA + ISLA Trial at NC A&T State University')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('For Monday, June 16 — Dr. Williams, Dr. Alston, Charlie Hopper, Chris Harrison')
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()  # spacer
    
    # Section 1: Opportunity
    doc.add_heading('1. The Opportunity', level=1)
    doc.add_paragraph(
        'The U.S. Department of Education will award approximately 6 grants averaging $3.75M each '
        '(48 months) to universities that use evidence-based, data-driven strategies to improve '
        'student retention and completion — and NC A&T already has a working platform that does exactly this.'
    )
    
    add_styled_table(doc,
        ['Detail', 'Value'],
        [
            ['Program', 'Postsecondary Student Success Grants (PSSG)'],
            ['Agency', 'U.S. Dept. of Education (administered by Dept. of Labor)'],
            ['Funding', '$45M total; ~$3.75M average per early phase award'],
            ['Number of Awards', '~6 early phase grants'],
            ['Project Period', '48 months'],
            ['Deadline', 'June 29, 2026'],
            ['Match', '10% (in-kind allowable)'],
        ]
    )
    
    # Section 2: Why NC A&T Wins
    doc.add_heading('2. Why NC A&T Wins This', level=1)
    
    doc.add_heading('Dr. Williams Makes This Application Elite', level=2)
    doc.add_paragraph(
        'Dr. Williams\' title — Vice Provost for Undergraduate Education and Student Success — '
        'mirrors the grant\'s name. She reports directly to the Provost. Reviewers will see '
        'instant institutional alignment.'
    )
    
    add_styled_table(doc,
        ['Without Dr. Williams', 'With Dr. Williams'],
        [
            ['CAES-only scope', 'University-wide scope under Vice Provost'],
            ['Department-level authority', 'Provost-level institutional authority'],
            ['Experiential learning in agriculture', 'All undergraduate pathways — 93 programs'],
            ['One college\'s student data', 'Institution-wide retention, completion, Pell data'],
            ['Competitive', 'Dominant'],
        ]
    )
    
    doc.add_heading('NC A&T\'s Competitive Advantages', level=2)
    bullets = [
        'Largest HBCU in the nation (~14,000 students)',
        '58% Pell Grant recipients — one of the strongest tiebreaker positions in the country',
        '1890 Land-Grant institution — unique federal relationship',
        '93 degree programs already mapped in EMMA — this is not a concept',
        'EMMA + ISLA already operational — AI coaching, career data, licensure tracking all working today',
        'Dr. Williams already endorses EMMA — "This is fantastic! I loove this!" (May 27, 2026)',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    # Section 3: The Intervention
    doc.add_heading('3. The Intervention: EMMA + ISLA', level=1)
    doc.add_paragraph(
        'EMMA (Experiential Major Mapping Assistant) is an AI-powered platform that maps every '
        'student\'s 4-year experiential journey from enrollment through career launch.'
    )
    doc.add_paragraph(
        'ISLA (Interactive Study & Licensure Assistant) maps coursework to professional exams '
        'and accreditation standards, showing students their competency progress in real time.'
    )
    doc.add_paragraph(
        'Together, they form a Curriculum-to-Credentials (C2C) Engine that:'
    )
    
    add_styled_table(doc,
        ['What It Does', 'How'],
        [
            ['Tracks experiential milestones', 'Interactive 4-year map with checkable milestones across 4 phases'],
            ['Provides AI career coaching', 'Google Gemini chat + neural voice (personalized per program)'],
            ['Shows live salary/employment data', 'Bureau of Labor Statistics API integration'],
            ['Maps coursework to licensure', 'ISLA validation bridge with animated progress gauges'],
            ['Connects students to opportunities', 'Handshake integration, SFRIC field projects, professional orgs'],
            ['Enables advisor oversight', 'Multi-role system (Student, Advisor, Admin)'],
            ['Works across any program', 'White-label architecture — 93 programs already seeded'],
        ]
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'What the grant would fund: A rigorous, controlled trial of EMMA + ISLA\'s impact on '
        'retention, credit accumulation, and completion — with an independent evaluation meeting '
        'What Works Clearinghouse standards.'
    )
    run.bold = True
    
    # Section 4: Team
    doc.add_heading('4. Proposed Team', level=1)
    add_styled_table(doc,
        ['Role', 'Person', 'What They Bring'],
        [
            ['Co-PI / Institutional Lead', 'Dr. Nakeshia Williams',
             'University-wide student success authority, Provost-level endorsement, institutional data access'],
            ['Co-PI / College Lead', 'Dr. Antoine Alston',
             'CAES Associate Dean, 25-year tenure, agricultural education expertise'],
            ['Co-PI / SFRIC Director', 'Charlie Hopper',
             'Experiential learning infrastructure, SFRIC field sites, XR/robotics, Martin Building CoLab'],
            ['Co-PI / Technology Director', 'Chris Harrison',
             'EMMA + ISLA platform architect, AI integration, data systems, Think! Design and Planning'],
            ['Project Director', 'To hire (Year 1)',
             'Day-to-day operations, student recruitment, data collection, reporting'],
            ['Independent Evaluator', 'To identify',
             'RCT or quasi-experimental design, WWC-compliant'],
        ]
    )
    
    # Section 5: Priorities
    doc.add_heading('5. Grant Priorities We Address', level=1)
    doc.add_paragraph('Applicants must address one evidence priority (AP1 or AP2) and one content priority (AP3-AP6).')
    doc.add_heading('Our Combination: AP1 + AP3', level=2)
    
    add_styled_table(doc,
        ['Priority', 'How We Meet It'],
        [
            ['AP1: Early Phase — Demonstrates a Rationale',
             'EMMA\'s logic model grounded in Kuh\'s High-Impact Practices framework. '
             'Pilot data from 93 programs provides rationale. Plan for WWC-compliant evaluation.'],
            ['AP3: Advancing AI to Support Student Success',
             'EMMA uses Google Gemini AI for personalized coaching, career guidance, and milestone-based advising. '
             'SFRIC adds XR/sensor/robotics technology layer.'],
            ['AP6: College-to-Career Pathways (secondary)',
             'EMMA\'s entire architecture IS a credential map. ISLA maps coursework to professional exams. '
             'BLS data shows career outcomes.'],
        ]
    )
    
    doc.add_heading('Competitive Preference Priority (10 bonus points)', level=2)
    doc.add_paragraph(
        '"Returning Education to the States" — Dr. Williams\' office can secure a partnership with '
        'the UNC System Office (Student Success Innovation Lab). NC A&T is a UNC System school — '
        'one letter of support qualifies. Worth 10 bonus points.'
    )
    
    # Section 6: Two-Grant Strategy
    doc.add_heading('6. Two-Grant Strategy', level=1)
    add_styled_table(doc,
        ['Grant', 'Amount', 'Deadline', 'What It Funds'],
        [
            ['NIFA Equipment Grants', 'Up to $500K', 'June 25', 'Martin Building CoLab — equipment, hardware (no match)'],
            ['FIPSE PSSG', 'Up to $3.75M', 'June 29', 'EMMA + ISLA trial — software, staffing, evaluation, student support'],
        ]
    )
    doc.add_paragraph(
        'Combined potential: $4.25M for CAES and NC A&T student success. '
        'NIFA equips the CoLab with hardware. PSSG funds the AI platform that connects students to '
        'that infrastructure and measures outcomes. Each makes the other stronger.'
    )
    
    # Section 7: Budget
    doc.add_heading('7. Budget Summary (PSSG)', level=1)
    add_styled_table(doc,
        ['Category', '48-Month Total', '%'],
        [
            ['Personnel (PIs, Project Director, coaches, GAs)', '$1,312,500', '35%'],
            ['Technology Development & Infrastructure', '$750,000', '20%'],
            ['Independent Evaluation', '$750,000', '20%'],
            ['Student Support Services', '$562,500', '15%'],
            ['Other Direct Costs', '$375,000', '10%'],
            ['TOTAL REQUEST', '$3,750,000', '100%'],
            ['10% Match (in-kind)', '$375,000', '—'],
        ]
    )
    
    # Section 8: Scope
    doc.add_heading('8. Proposed Scope', level=1)
    add_styled_table(doc,
        ['Phase', 'Timeline', 'Scope', 'Students'],
        [
            ['Pilot', 'Years 1–2', 'CAES (17 programs) + select programs from 1–2 other colleges', '~3,000'],
            ['Scale', 'Years 3–4', 'Expand to 4+ colleges, 50+ programs', '~6,000+'],
        ]
    )
    
    # Section 9: Decisions for Monday
    doc.add_heading('9. Decisions Needed Monday, June 16', level=1)
    
    doc.add_heading('For Dr. Williams:', level=2)
    decisions_w = [
        'Are you willing to serve as Co-PI / Institutional Lead?',
        'Can your office provide institutional retention/completion data for the narrative?',
        'Can you initiate contact with UNC System Office for a letter of support (10 bonus points)?',
        'Will the Office of Sponsored Programs support a June 29 submission?',
    ]
    for d in decisions_w:
        doc.add_paragraph(d, style='List Number')
    
    doc.add_heading('For Dr. Alston:', level=2)
    decisions_a = [
        'Are you willing to serve as Co-PI / College Lead?',
        'Do you have published research we can cite as evidence?',
        'Do you have existing relationships with evaluation firms (RTI, SERVE Center, MDRC)?',
    ]
    for d in decisions_a:
        doc.add_paragraph(d, style='List Number')
    
    doc.add_heading('For All:', level=2)
    decisions_all = [
        'Scope confirmation — CAES pilot + selected programs from other colleges?',
        'Evaluation partner — who do we approach first?',
        'Division of labor for the next 13 days?',
    ]
    for d in decisions_all:
        doc.add_paragraph(d, style='List Number')
    
    # Section 10: Why this grant was written for us
    doc.add_heading('10. Why This Grant Was Written for Us', level=1)
    p = doc.add_paragraph()
    run = p.add_run('The PSSG NOFO specifically encourages:')
    run.italic = True
    
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.5)
    run = quote.add_run(
        '"projects that leverage AI technology to improve developmental education; projects designed '
        'to serve as on-ramps to career pathways; projects designed to lead to credentials; and other '
        'evidence-based strategies designed to improve credit accumulation, college retention and completion"'
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    run = p.add_run(
        'That is a literal description of EMMA + ISLA. NC A&T doesn\'t need to build anything new — '
        'we need to trial what\'s already built and prove it works. That\'s exactly what early phase grants are for.'
    )
    run.bold = True
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Prepared by Chris Harrison | Think! Design and Planning, LLC | June 14, 2026')
    run.font.size = Pt(8)
    run.font.italic = True
    
    # Save
    filepath = os.path.join(OUTPUT_DIR, 'EMMA_PSSG_Strategy_Brief.docx')
    doc.save(filepath)
    print(f"✅ Strategy Brief saved: {filepath}")
    return filepath


def create_project_abstract():
    """Create the Project Abstract Word document (1 double-spaced page)."""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # Double-spaced
    style.paragraph_format.space_after = Pt(0)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Project Abstract')
    run.bold = True
    run.font.size = Pt(12)
    
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run(
        'EMMA: An AI-Powered Experiential Major Mapping Assistant '
        'for Postsecondary Student Success'
    )
    run.bold = True
    run.font.size = Pt(12)
    
    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title3.add_run('North Carolina Agricultural and Technical State University')
    run.font.size = Pt(12)
    run.italic = True
    
    # Priorities line
    p = doc.add_paragraph()
    run = p.add_run('Absolute Priorities: ')
    run.bold = True
    p.add_run('AP1 (Early Phase — Demonstrates a Rationale); AP3 (Advancing AI to Support Student Success)')
    
    p = doc.add_paragraph()
    run = p.add_run('Competitive Preference Priority: ')
    run.bold = True
    p.add_run('Returning Education to the States')
    
    # Body
    doc.add_paragraph(
        'North Carolina Agricultural and Technical State University (NC A&T), the nation\'s largest '
        'Historically Black University serving approximately 14,000 students — 58% of whom are Pell '
        'Grant recipients — proposes a 48-month early phase project to implement and evaluate EMMA '
        '(Experiential Major Mapping Assistant), an AI-powered Curriculum-to-Credentials platform '
        'designed to improve student retention, credit accumulation, and postsecondary completion.'
    )
    
    doc.add_paragraph(
        'EMMA, developed by Think! Design and Planning, LLC in partnership with NC A&T, is an '
        'operational technology platform that maps every student\'s four-year experiential journey '
        'across four developmental phases — Explore, Engage, Develop, and Launch — connecting '
        'coursework to career outcomes through interactive milestone tracking, AI-powered coaching '
        'via Google Gemini, live Bureau of Labor Statistics employment data, and professional '
        'licensure validation through its companion module, ISLA (Interactive Study & Licensure '
        'Assistant). The platform currently supports 93 degree programs across all eight NC A&T '
        'colleges with a white-label architecture enabling deployment to any institution.'
    )
    
    doc.add_paragraph(
        'The proposed project will conduct a rigorous trial of EMMA + ISLA beginning with a pilot '
        'in the College of Agriculture and Environmental Sciences (17 programs, ~3,000 students) '
        'in Years 1-2, scaling to 4+ colleges and 6,000+ students in Years 3-4. The project team '
        'includes Dr. Nakeshia Williams (Vice Provost for Undergraduate Education and Student '
        'Success), Dr. Antoine Alston (Associate Dean for Academic Studies, CAES), Charlie Hopper '
        '(Director, Small Farm Research and Innovation Center), and W. Christopher Harrison '
        '(EMMA platform architect). An independent evaluator will design and implement a study '
        'meeting What Works Clearinghouse standards to assess EMMA\'s impact on persistence, credit '
        'accumulation, retention, completion, and post-completion career outcomes.'
    )
    
    doc.add_paragraph(
        'The project addresses AP3 (Advancing AI) by deploying EMMA\'s AI coaching engine to '
        'provide personalized, milestone-based guidance to students at critical junctures in their '
        'academic journey. Through a state-level partnership with the University of North Carolina '
        'System Office, the project will generate replicable evidence for scaling AI-powered '
        'experiential learning platforms across the UNC System and beyond.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Proposed Outcomes: ')
    run.bold = True
    p.add_run(
        'Increase first-to-second year retention by 5 percentage points; increase 4-year '
        'completion rate by 8 percentage points; achieve 75% EMMA engagement rate among enrolled '
        'students; and produce WWC-eligible evidence on the effectiveness of AI-powered experiential '
        'mapping for postsecondary student success at HBCUs.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Project Partners: ')
    run.bold = True
    p.add_run(
        'NC A&T State University (OPEID: 00256200), Think! Design and Planning, LLC, '
        'UNC System Office, Small Farm Research and Innovation Center.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Requested Funding: ')
    run.bold = True
    p.add_run('$3,750,000 over 48 months.')
    
    # Save
    filepath = os.path.join(OUTPUT_DIR, 'EMMA_PSSG_Project_Abstract.docx')
    doc.save(filepath)
    print(f"✅ Project Abstract saved: {filepath}")
    return filepath


if __name__ == '__main__':
    print("Generating FIPSE PSSG Grant Documents...")
    print()
    create_strategy_brief()
    create_project_abstract()
    print()
    print("📁 Both documents saved to Desktop. Upload to Charlie's Google Drive:")
    print("   https://drive.google.com/drive/folders/1ECCpji7iYWiM10SlQsQcRWSMxomH6hkP")
