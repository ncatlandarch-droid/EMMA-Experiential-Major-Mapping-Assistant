import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor

doc = Document(r'C:\Users\Chris\Downloads\BS in Animal Science.docx')

# Get all runs with formatting
print('=== DETAILED PARAGRAPHS WITH FORMATTING ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text and len(para.runs) == 0:
        continue
    align = para.alignment
    style = para.style.name if para.style else 'None'
    print(f'\nP{i} [style={style}, align={align}]')
    for j, run in enumerate(para.runs):
        rt = run.text.strip()
        if not rt:
            continue
        font = run.font
        print(f'  R{j}: "{rt[:120]}"')
        print(f'       bold={font.bold}, italic={font.italic}, size={font.size}, color={font.color.rgb if font.color and font.color.rgb else None}, name={font.name}')

# Table cell formatting details
print('\n=== TABLE CELL SHADING ===')
table = doc.tables[0]
for ri, row in enumerate(table.rows):
    for ci, cell in enumerate(row.cells):
        text = cell.text.strip()[:60]
        tc = cell._tc
        tcPr = tc.tcPr
        shading = None
        if tcPr is not None:
            shd = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd is not None:
                shading = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        print(f'  Cell[{ri},{ci}]: bg={shading}, "{text}"')

# Check for images
print('\n=== IMAGES ===')
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        print(f'  Image: {rel.target_ref}')
